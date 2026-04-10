from pathlib import Path
from docx import Document
from docx.shared import Inches

BASE = Path(__file__).resolve().parent
INPUT = BASE / 'SOLUTION_REPORT.md'
OUTPUT = BASE / 'SOLUTION_REPORT.docx'
IMAGES = {
    '## 3. Diagrama de funcionamiento': BASE / 'ARCHITECTURE.png',
    '## 5. Diagrama SQL de la base de datos': BASE / 'SQL_SCHEMA.png',
}

doc = Document()
with INPUT.open('r', encoding='utf-8') as f:
    lines = [line.rstrip('\n') for line in f]

in_code = False
code_lines = []
for idx, line in enumerate(lines):
    if line.startswith('```'):
        if in_code:
            paragraph = doc.add_paragraph()
            paragraph.style = 'NoSpacing'
            run = paragraph.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            in_code = False
            code_lines = []
        else:
            in_code = True
        continue
    if in_code:
        code_lines.append(line)
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:], level=0)
        continue
    if line.startswith('## '):
        heading_text = line[3:]
        doc.add_heading(heading_text, level=1)
        key = line
        if key in IMAGES and IMAGES[key].exists():
            doc.add_picture(str(IMAGES[key]), width=Inches(6))
        continue
    if line.startswith('### '):
        doc.add_heading(line[4:], level=2)
        continue
    if line.startswith('- '):
        paragraph = doc.add_paragraph(line[2:], style='List Bullet')
        continue
    if line.strip() == '':
        doc.add_paragraph('')
        continue
    doc.add_paragraph(line)

if in_code and code_lines:
    paragraph = doc.add_paragraph()
    paragraph.style = 'NoSpacing'
    run = paragraph.add_run('\n'.join(code_lines))
    run.font.name = 'Courier New'

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f'Word document generated: {OUTPUT}')
